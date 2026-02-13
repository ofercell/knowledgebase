"""DOCX document processor with Hebrew text support."""

from pathlib import Path
from typing import List, Dict, Any

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table
except ImportError:
    Document = None

from .base import DocumentProcessor, DocumentChunk
from ..config import config


class DOCXProcessor(DocumentProcessor):
    """Process DOCX documents with support for Hebrew text."""
    
    def __init__(self):
        """Initialize DOCX processor."""
        if Document is None:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install it with: pip install python-docx"
            )
    
    def can_process(self, file_path: Path) -> bool:
        """Check if file is a DOCX."""
        return file_path.suffix.lower() in ['.docx', '.doc']
    
    def process(self, file_path: Path) -> List[DocumentChunk]:
        """
        Process DOCX file and extract text content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        metadata = self.extract_metadata(file_path)
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)
            
            # Combine all text
            full_text = "\n".join(text_content)
            
            # Split into chunks
            text_chunks = self._split_text(full_text)
            
            for idx, chunk_text in enumerate(text_chunks):
                chunk = DocumentChunk(
                    content=chunk_text,
                    metadata={
                        **metadata,
                        "chunk_type": "text",
                        "supports_hebrew": True
                    },
                    chunk_index=idx
                )
                chunks.append(chunk)
            
            # Extract images (stored as relationships in DOCX)
            # Note: Image data is embedded in the document
            # For simplicity, we'll note their presence in metadata
            image_count = len(doc.inline_shapes)
            if image_count > 0:
                metadata["image_count"] = image_count
                
        except Exception as e:
            raise RuntimeError(f"Error processing DOCX {file_path}: {str(e)}")
        
        return chunks
    
    def _split_text(self, text: str) -> List[str]:
        """
        Split text into chunks based on configuration.
        Handles Hebrew text properly.
        
        Args:
            text: Text to split
            
        Returns:
            List of text chunks
        """
        chunk_size = config.CHUNK_SIZE
        chunk_overlap = config.CHUNK_OVERLAP
        
        # Validate chunk parameters to prevent infinite loop
        if chunk_overlap >= chunk_size:
            chunk_overlap = chunk_size // 2
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - chunk_overlap
        
        return chunks
